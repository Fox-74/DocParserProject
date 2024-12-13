from docx import Document
import pandas as pd
import mysql.connector
import re
import os

# Настройки базы данных
DB_CONFIG = {
    'host': 'localhost',
    'user': 'root',
    'password': 'password',
    'database': 'commands_db'
}


# Функция для извлечения данных из Word-файла
def extract_tables_from_word(file_path, section_title):
    document = Document(file_path)
    tables = []
    is_target_section = False

    for paragraph in document.paragraphs:
        if section_title in paragraph.text:
            is_target_section = True
        elif is_target_section and paragraph.style.name.startswith('Heading'):
            break  # Выходим, если следующий раздел начался

    if is_target_section:
        for table in document.tables:
            tables.append([[cell.text.strip() for cell in row.cells] for row in table.rows])

    return tables


# Функция для обработки ссылок на Excel
def process_excel_links(table, excel_folder):
    data = []
    for row in table:
        for cell in row:
            match = re.search(r'(.*\.xlsx)', cell)
            if match:
                excel_path = os.path.join(excel_folder, match.group(1))
                if os.path.exists(excel_path):
                    excel_data = pd.read_excel(excel_path)
                    data.append(excel_data)
    return data


# Функция для записи данных в MySQL
def write_to_mysql(data, table_name):
    connection = mysql.connector.connect(**DB_CONFIG)
    cursor = connection.cursor()

    for index, row in data.iterrows():
        placeholders = ', '.join(['%s'] * len(row))
        columns = ', '.join(row.index)
        sql = f"INSERT INTO {table_name} ({columns}) VALUES ({placeholders})"
        cursor.execute(sql, tuple(row))

    connection.commit()
    cursor.close()
    connection.close()


# Основная программа
def main():
    word_file = "path_to_word_file.docx"
    excel_folder = "path_to_excel_files"
    section_title = "Название раздела"
    mysql_table = "commands"

    print("Извлечение таблиц из Word...")
    tables = extract_tables_from_word(word_file, section_title)

    for table in tables:
        print("Обработка ссылок на Excel...")
        excel_data = process_excel_links(table, excel_folder)

        for df in excel_data:
            print("Запись данных в MySQL...")
            write_to_mysql(df, mysql_table)

    print("Процесс завершен!")


if __name__ == "__main__":
    main()